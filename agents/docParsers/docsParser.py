import logging
import sys
from io import BytesIO
from pathlib import Path

import requests
from docx import Document

from docParsers.ollamaWorker import FastSummarizer
from sse.event_bus import event_bus
from utils.logger.AgentLogger import quickLog
from utils.task_scheduler import scheduler

BASE_DIR = Path(__file__).resolve().parent

# ====================== CLEAN LOGGER SETUP ======================
LOG_FORMAT = "%(asctime)s  %(levelname)-8s  %(message)s"
logging.basicConfig(
    level=logging.INFO,
    format=LOG_FORMAT,
    datefmt="%H:%M:%S",
    handlers=[logging.StreamHandler(sys.stdout)],
)

# Suppress noisy libraries
logging.getLogger("urllib3").setLevel(logging.WARNING)
logging.getLogger("docx").setLevel(logging.WARNING)

logger = logging.getLogger("DOCXParser")


def extract_docx_to_md(docx_path):
    """Extract content from .docx - supports both local path and URL.
    Raises exception on failure so caller can handle it."""

    is_url = isinstance(docx_path, str) and docx_path.startswith(
        ("http://", "https://")
    )

    # ---------- Handle URL or Local File ----------
    if is_url:
        logger.info(f"Downloading DOCX from URL: {docx_path}")
        try:
            response = requests.get(docx_path, timeout=60)
            response.raise_for_status()
            file_stream = BytesIO(response.content)
            logger.info("DOCX download successful")
        except requests.exceptions.HTTPError as e:
            logger.error(f"HTTP Error downloading {docx_path}: {e}")
            raise
        except Exception as e:
            logger.error(f"Failed to download DOCX from {docx_path}: {e}")
            raise
    else:
        logger.info(f"Opening local DOCX: {docx_path}")
        file_stream = open(docx_path, "rb")

    try:
        doc = Document(file_stream)

        md = []

        # -------- TEXT + HEADINGS --------
        logger.info("Extracting content from DOCX")
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name.lower()
            if "heading" in style:
                level = "".join(filter(str.isdigit, style))
                level = int(level) if level else 1
                md.append(f"{'#' * level} {text}")
            else:
                md.append(text)

        # -------- TABLES --------
        for table in doc.tables:
            rows = []
            logger.info("Extracting tables from DOCX")

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if not rows:
                continue

            header = rows[0]
            md.append("\n| " + " | ".join(header) + " |")
            md.append("| " + " | ".join(["---"] * len(header)) + " |")

            for row in rows[1:]:
                md.append("| " + " | ".join(row) + " |")

            md.append("\n")

        return "\n".join(md)

    except Exception as e:
        logger.error(f"Failed to parse DOCX {docx_path}: {e}")
        raise
    finally:
        # Close file only if we opened it (local file)
        if not is_url and "file_stream" in locals():
            file_stream.close()


async def extract_text_summarized(docx_path, summarizer_url):
    try:
        text = extract_docx_to_md(docx_path)
        await scheduler.schedule(
            quickLog,
            params={
                "message": f"Extracted {docx_path} now summarizing it!",
                "level": "success",
                "urgency": "none",
                "module": ["UTILS"],
            },
        )
        logger.info(f"Extracted DOCX content, now summarizing: {docx_path}")
        return FastSummarizer(base_url=summarizer_url).summarize(text)

    except Exception as e:
        logger.error(f"Failed to process {docx_path} for summarization: {e}")
        return f"ERROR: Failed to extract or summarize {docx_path}\n{str(e)}"


async def bulk_extract_text_summarized(docx_paths, summarizer_url):
    results = {}
    await event_bus.broadcast(message={"msg": "Reading & Summarizing DOCX Files!"})

    for docx_path in docx_paths:
        try:
            results[docx_path] = await extract_text_summarized(
                docx_path, summarizer_url
            )
        except Exception as e:
            logger.warning(f"Skipped {docx_path} due to error: {e}")
            results[docx_path] = f"ERROR: Could not process this file. {str(e)}"

    return results


# ====================== TEST ======================
# async def _test():
#     docx_list = [
#         r"https://sample-files.com/downloads/documents/docx/sample-files.com-basic-text.docx",  # Good
#         "https://www.orimi.com/pdf-test.pdf",                                               # Bad (wrong extension + 403)
#         # r"D:\path\to\your\local\file.docx"   # Local files still work perfectly
#     ]

#     result = await bulk_extract_text_summarized(docx_list, "http://localhost:11434/api/generate")

#     print("\n=== FINAL RESULT ===")
#     for path, content in result.items():
#         print(f"\n--- {path} ---")
#         preview = str(content)[:600] + "..." if len(str(content)) > 600 else content
#         print(preview)


# if __name__ == "__main__":
#     import asyncio
#     asyncio.run(_test())
